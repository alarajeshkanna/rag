"""
Document loader module for different document types
"""
import requests
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
import logging
from pathlib import Path
import PyPDF2
from docx import Document
import io
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import tempfile
import os

from config import LoaderType, KGConfig

logger = logging.getLogger(__name__)


class DocumentLoader(ABC):
    """Abstract base class for document loaders"""
    
    @abstractmethod
    def load(self, source: str, **kwargs) -> List[Dict[str, Any]]:
        """Load documents from source"""
        pass


class URLLoader(DocumentLoader):
    """Load documents from URLs"""
    
    def __init__(self, config: KGConfig):
        self.config = config
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
    
    def load(self, source: str, **kwargs) -> List[Dict[str, Any]]:
        """Load content from URL"""
        try:
            response = self.session.get(source, timeout=30)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # Extract metadata
            title = soup.find('title')
            title_text = title.get_text().strip() if title else urlparse(source).netloc
            
            meta_description = soup.find('meta', attrs={'name': 'description'})
            description = meta_description.get('content', '') if meta_description else ''
            
            return [{
                'content': text,
                'metadata': {
                    'source': source,
                    'title': title_text,
                    'description': description,
                    'type': 'url',
                    'length': len(text)
                }
            }]
            
        except Exception as e:
            logger.error(f"Error loading URL {source}: {str(e)}")
            raise


class PDFLoader(DocumentLoader):
    """Load documents from PDF files"""
    
    def __init__(self, config: KGConfig):
        self.config = config
    
    def load(self, source: str, **kwargs) -> List[Dict[str, Any]]:
        """Load content from PDF file"""
        try:
            documents = []
            
            with open(source, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text += page_text + "\n"
                    
                    # Create document for each page if requested
                    if kwargs.get('page_wise', False):
                        documents.append({
                            'content': page_text,
                            'metadata': {
                                'source': source,
                                'page': page_num + 1,
                                'type': 'pdf',
                                'total_pages': len(pdf_reader.pages),
                                'length': len(page_text)
                            }
                        })
                
                # Create single document with full content if not page-wise
                if not kwargs.get('page_wise', False):
                    # Extract metadata from PDF
                    metadata = pdf_reader.metadata if pdf_reader.metadata else {}
                    
                    documents.append({
                        'content': full_text,
                        'metadata': {
                            'source': source,
                            'title': metadata.get('/Title', Path(source).stem),
                            'author': metadata.get('/Author', ''),
                            'subject': metadata.get('/Subject', ''),
                            'type': 'pdf',
                            'total_pages': len(pdf_reader.pages),
                            'length': len(full_text)
                        }
                    })
            
            return documents
            
        except Exception as e:
            logger.error(f"Error loading PDF {source}: {str(e)}")
            raise


class DOCXLoader(DocumentLoader):
    """Load documents from DOCX files"""
    
    def __init__(self, config: KGConfig):
        self.config = config
    
    def load(self, source: str, **kwargs) -> List[Dict[str, Any]]:
        """Load content from DOCX file"""
        try:
            doc = Document(source)
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    full_text += row_text + "\n"
            
            # Extract core properties
            props = doc.core_properties
            
            documents = []
            
            # Create document per paragraph if requested
            if kwargs.get('paragraph_wise', False):
                for i, para_text in enumerate(paragraphs):
                    documents.append({
                        'content': para_text,
                        'metadata': {
                            'source': source,
                            'paragraph': i + 1,
                            'type': 'docx',
                            'title': props.title or Path(source).stem,
                            'author': props.author or '',
                            'length': len(para_text)
                        }
                    })
            else:
                # Single document with full content
                documents.append({
                    'content': full_text,
                    'metadata': {
                        'source': source,
                        'title': props.title or Path(source).stem,
                        'author': props.author or '',
                        'subject': props.subject or '',
                        'type': 'docx',
                        'created': props.created.isoformat() if props.created else '',
                        'modified': props.modified.isoformat() if props.modified else '',
                        'length': len(full_text)
                    }
                })
            
            return documents
            
        except Exception as e:
            logger.error(f"Error loading DOCX {source}: {str(e)}")
            raise


class DocumentLoaderFactory:
    """Factory class for creating document loaders"""
    
    def __init__(self, config: KGConfig):
        self.config = config
        self._loaders = {
            LoaderType.URL: URLLoader(config),
            LoaderType.PDF: PDFLoader(config),
            LoaderType.DOCX: DOCXLoader(config)
        }
    
    def get_loader(self, loader_type: LoaderType) -> DocumentLoader:
        """Get appropriate loader for document type"""
        if loader_type not in self._loaders:
            raise ValueError(f"Unsupported loader type: {loader_type}")
        return self._loaders[loader_type]
    
    def load_document(self, source: str, loader_type: LoaderType, **kwargs) -> List[Dict[str, Any]]:
        """Load document using appropriate loader"""
        loader = self.get_loader(loader_type)
        return loader.load(source, **kwargs)
    
    def auto_detect_and_load(self, source: str, **kwargs) -> List[Dict[str, Any]]:
        """Auto-detect document type and load"""
        if source.startswith(('http://', 'https://')):
            return self.load_document(source, LoaderType.URL, **kwargs)
        elif source.lower().endswith('.pdf'):
            return self.load_document(source, LoaderType.PDF, **kwargs)
        elif source.lower().endswith(('.docx', '.doc')):
            return self.load_document(source, LoaderType.DOCX, **kwargs)
        else:
            raise ValueError(f"Cannot auto-detect document type for: {source}")
